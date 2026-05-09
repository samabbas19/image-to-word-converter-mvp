import base64
import mimetypes
import os
import re
import uuid
from dataclasses import asdict, dataclass, field
from datetime import datetime
from typing import Dict, List, Optional, Tuple

import cv2
import numpy as np
from docx import Document
from docx.shared import Inches, Pt
from dotenv import load_dotenv
from groq import Groq
from PIL import Image

from diagrams import crop_diagrams, extract_diagrams, parse_diagram_bounds


load_dotenv()

OCR_PROMPT = (
    "You are an OCR and document reconstruction engine.\n\n"
    "Extract all visible handwritten and printed content from the image. Preserve reading order, "
    "headings, bullets, equations, arrows, columns, tables, and diagram labels as plain text.\n\n"
    "Rules:\n"
    "- Output only the extracted content.\n"
    "- Do not summarize, explain, correct spelling, or add new information.\n"
    "- Use [illegible] for unclear text.\n"
    "- Preserve line breaks and indentation where possible.\n"
    "- For two-column handwritten notes, transcribe the left column first, then the right column.\n"
)

AGENTIC_REQUIREMENTS = [
    "Perception: inspect uploaded page quality before conversion.",
    "Interpretation: classify image risks into a quality status.",
    "Decision-making: select preprocessing, review policy, and diagram strategy.",
    "Action: call OCR, diagram extraction, cropping, and DOCX generation tools.",
    "Memory/context: retain run-level outcomes without permanent user-note storage.",
    "Human-in-the-loop: keep the user responsible for final acceptance.",
    "Ethical design: disclose OCR uncertainty, privacy limits, and user control.",
    "Safety: log decisions, protect API keys, and flag risky pages.",
]

SAFETY_CONTROLS = [
    "Groq API key is read from environment variables.",
    "No OCR or diagram API call runs during module import.",
    "Uploaded files are limited to common image formats by the Streamlit UI.",
    "Low-quality pages are flagged for human review.",
    "Agent decisions are written to the UI trace, while the DOCX contains only transformed source content.",
    "Run memory is temporary and does not retain user notes by default.",
]

RISK_REGISTER = [
    {
        "risk": "OCR misreads handwriting or formulas",
        "impact": "Wrong academic notes",
        "mitigation": "Quality gate, trace, text preview, and DOCX review",
    },
    {
        "risk": "Private notes are sent to a third-party OCR provider",
        "impact": "Privacy exposure",
        "mitigation": "Explicit disclosure and no long-term local memory by default",
    },
    {
        "risk": "API key leakage",
        "impact": "Unauthorized API usage",
        "mitigation": "Environment variables and .gitignore protection",
    },
    {
        "risk": "Copyrighted material is uploaded without permission",
        "impact": "IPR and acceptable-use violation",
        "mitigation": "User ownership notice and acceptable-use reminder",
    },
    {
        "risk": "Over-automation",
        "impact": "User trusts a draft as final",
        "mitigation": "Semi-autonomous positioning and human review flag",
    },
]

EXTERNAL_TOOLS = [
    "Groq vision OCR",
    "Groq vision layout analysis",
    "OpenCV",
    "Pillow",
    "python-docx",
]

DEFAULT_GROQ_VISION_MODEL = "meta-llama/llama-4-scout-17b-16e-instruct"
RISK_RANK = {"Low": 0, "Medium": 1, "High": 2}


@dataclass
class ImageQualityReport:
    filename: str
    width: int
    height: int
    brightness: float
    contrast: float
    blur_score: float
    status: str
    issues: List[str] = field(default_factory=list)
    recommendations: List[str] = field(default_factory=list)


@dataclass
class AgentDecision:
    autonomy_level: str
    preprocessing_applied: List[str]
    diagram_strategy: str
    human_review_required: bool
    rationale: List[str]
    risk_level: str = "Low"
    review_policy: str = "quality_gate"
    review_checkpoint: str = "User reviews the trace and DOCX before relying on the output."
    safety_controls: List[str] = field(default_factory=list)
    external_tools: List[str] = field(default_factory=list)


@dataclass
class PageConversionResult:
    source_path: str
    processed_path: str
    extracted_text: str
    diagram_data: List[Dict[str, object]]
    quality: ImageQualityReport
    decision: AgentDecision
    trace: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    status: str = "completed"


def get_groq_client() -> Groq:
    api_key = os.getenv("GROQ_API_KEY") or os.getenv("GROK_API_KEY")
    if not api_key:
        raise RuntimeError(
            "Missing Groq API key. Set GROQ_API_KEY in .env "
            "(GROK_API_KEY is also supported for backward compatibility)."
        )
    return Groq(api_key=api_key)


def get_groq_vision_model() -> str:
    return os.getenv("GROQ_VISION_MODEL", DEFAULT_GROQ_VISION_MODEL)


def image_bytes_to_data_url(image_bytes: bytes, filename: str = "image.png") -> str:
    mime_type, _ = mimetypes.guess_type(filename)
    if mime_type is None:
        mime_type = "image/png"

    encoded = base64.b64encode(image_bytes).decode("utf-8")
    return "data:{0};base64,{1}".format(mime_type, encoded)


def _highest_risk(*levels: str) -> str:
    return max((level for level in levels if level), key=lambda item: RISK_RANK.get(item, 0), default="Low")


def _quality_risk_level(status: str) -> str:
    if status == "Human review required":
        return "High"
    if status == "Review recommended":
        return "Medium"
    return "Low"


def _append_failure(decision: AgentDecision, message: str, risk_level: str = "High") -> None:
    decision.human_review_required = True
    decision.risk_level = _highest_risk(decision.risk_level, risk_level)
    decision.rationale.append(message)


def extract_text_from_image(image_bytes: bytes, filename: str) -> str:
    image_url = image_bytes_to_data_url(image_bytes, filename)
    request_payload = {
        "model": get_groq_vision_model(),
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": OCR_PROMPT},
                    {"type": "image_url", "image_url": {"url": image_url}},
                ],
            }
        ],
        "temperature": 0,
    }
    try:
        completion = get_groq_client().chat.completions.create(**request_payload, max_completion_tokens=4096)
    except TypeError:
        completion = get_groq_client().chat.completions.create(**request_payload, max_tokens=4096)
    return completion.choices[0].message.content or ""


def assess_image_quality(image_path: str) -> ImageQualityReport:
    with Image.open(image_path) as image:
        width, height = image.size

    cv_image = cv2.imread(image_path)
    if cv_image is None:
        raise ValueError("Could not read image: {0}".format(image_path))

    gray = cv2.cvtColor(cv_image, cv2.COLOR_BGR2GRAY)
    brightness = float(np.mean(gray))
    contrast = float(np.std(gray))
    blur_score = float(cv2.Laplacian(gray, cv2.CV_64F).var())

    issues = []
    recommendations = []

    if min(width, height) < 900:
        issues.append("Resolution is low for handwritten OCR.")
        recommendations.append("Use a higher-resolution scan or photo when possible.")
    if blur_score < 80:
        issues.append("Image may be blurry.")
        recommendations.append("Retake the photo with the page flat and camera steady.")
    if contrast < 35:
        issues.append("Text contrast is low.")
        recommendations.append("Increase contrast or use brighter, even lighting.")
    if brightness < 80 or brightness > 210:
        issues.append("Lighting may be uneven or too bright/dark.")
        recommendations.append("Use neutral lighting and avoid shadows or glare.")

    if not issues:
        status = "Ready"
    elif len(issues) <= 2:
        status = "Review recommended"
    else:
        status = "Human review required"

    return ImageQualityReport(
        filename=os.path.basename(image_path),
        width=width,
        height=height,
        brightness=round(brightness, 2),
        contrast=round(contrast, 2),
        blur_score=round(blur_score, 2),
        status=status,
        issues=issues,
        recommendations=recommendations,
    )


def prepare_image_for_ocr(
    image_path: str, quality: ImageQualityReport, output_dir: str
) -> Tuple[str, List[str]]:
    os.makedirs(output_dir, exist_ok=True)
    image = cv2.imread(image_path)
    if image is None:
        raise ValueError("Could not read image: {0}".format(image_path))

    operations = []
    processed = image.copy()

    if quality.contrast < 45 or quality.brightness < 95 or quality.brightness > 195:
        lab = cv2.cvtColor(processed, cv2.COLOR_BGR2LAB)
        l_channel, a_channel, b_channel = cv2.split(lab)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced_l = clahe.apply(l_channel)
        processed = cv2.cvtColor(cv2.merge((enhanced_l, a_channel, b_channel)), cv2.COLOR_LAB2BGR)
        operations.append("contrast normalization")

    if quality.blur_score < 140:
        kernel = np.array([[0, -1, 0], [-1, 5, -1], [0, -1, 0]])
        processed = cv2.filter2D(processed, -1, kernel)
        operations.append("light sharpening")

    if not operations:
        return image_path, operations

    filename = "{0}_prepared.png".format(uuid.uuid4().hex[:12])
    output_path = os.path.join(output_dir, filename)
    cv2.imwrite(output_path, processed)
    return output_path, operations


def detect_formatting(text: str) -> List[Dict[str, object]]:
    lines = text.split("\n")
    formatted_content = []

    for line in lines:
        stripped = line.strip()
        line_data = {
            "text": line,
            "is_heading": False,
            "is_bullet": False,
            "is_numbered": False,
            "indent_level": len(line) - len(line.lstrip()),
        }

        if stripped.startswith("#") or (
            stripped.endswith(":") and len(stripped) < 80 and not re.match(r"^\d+[\.\)]", stripped)
        ):
            line_data["is_heading"] = True
        elif stripped.isupper() and 3 <= len(stripped) < 80:
            line_data["is_heading"] = True

        if re.match(r"^\s*([*\\-]|o|•)\s+", line):
            line_data["is_bullet"] = True
            line_data["is_heading"] = False

        if re.match(r"^\s*\d+[\.\)]\s+", line):
            line_data["is_numbered"] = True
            line_data["is_heading"] = False

        formatted_content.append(line_data)

    return formatted_content


def _set_document_defaults(doc: Document) -> None:
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)


def _add_text_with_inline_diagrams(
    doc: Document,
    text: str,
    diagram_data: List[Dict[str, object]],
) -> None:
    formatted_content = detect_formatting(text)
    total_lines = max(1, len(formatted_content))
    diagrams = sorted(diagram_data, key=lambda item: item["y_min"])
    diagram_index = 0

    for line_index, line_data in enumerate(formatted_content):
        line_y_position = (line_index / total_lines) * 100

        while diagram_index < len(diagrams):
            diagram = diagrams[diagram_index]
            if diagram["y_min"] <= line_y_position + 4:
                diagram_path = str(diagram["path"])
                if os.path.exists(diagram_path):
                    doc.add_picture(diagram_path, width=Inches(5.6))
                    doc.add_paragraph()
                diagram_index += 1
            else:
                break

        text_content = str(line_data["text"]).strip()
        if not text_content:
            doc.add_paragraph()
            continue

        paragraph = doc.add_paragraph()
        indent_level = int(line_data["indent_level"])

        if line_data["is_heading"]:
            paragraph.style = "Heading 2"
            run = paragraph.add_run(text_content.lstrip("#").strip())
            run.bold = True
        elif line_data["is_bullet"]:
            paragraph.style = "List Bullet"
            clean_text = re.sub(r"^\s*([*\\-]|o|•)\s+", "", text_content)
            paragraph.add_run(clean_text)
        elif line_data["is_numbered"]:
            paragraph.style = "List Number"
            clean_text = re.sub(r"^\s*\d+[\.\)]\s+", "", text_content)
            paragraph.add_run(clean_text)
        else:
            paragraph.add_run(text_content)

        if indent_level > 0:
            paragraph.paragraph_format.left_indent = Inches(min(indent_level * 0.04, 1.2))

    while diagram_index < len(diagrams):
        diagram = diagrams[diagram_index]
        diagram_path = str(diagram["path"])
        if os.path.exists(diagram_path):
            doc.add_picture(diagram_path, width=Inches(5.6))
            doc.add_paragraph()
        diagram_index += 1


def create_agentic_docx(
    results: List[PageConversionResult],
    output_path: str,
    memory: Optional[Dict[str, object]] = None,
) -> str:
    doc = Document()
    _set_document_defaults(doc)

    for page_index, result in enumerate(results, start=1):
        if page_index > 1:
            doc.add_page_break()

        _add_text_with_inline_diagrams(doc, result.extracted_text, result.diagram_data)

    doc.save(output_path)
    return output_path


class DocumentConversionAgent:
    def __init__(
        self,
        work_dir: str = "temp",
        diagram_dir: str = "cropped_diagrams",
        review_policy: str = "quality_gate",
    ) -> None:
        self.work_dir = work_dir
        self.diagram_dir = diagram_dir
        self.review_policy = review_policy
        os.makedirs(self.work_dir, exist_ok=True)
        os.makedirs(self.diagram_dir, exist_ok=True)

    def observe(self, image_path: str) -> ImageQualityReport:
        return assess_image_quality(image_path)

    def decide(self, quality: ImageQualityReport, preprocessing: List[str]) -> AgentDecision:
        human_review_required = quality.status != "Ready" or self.review_policy == "review_all"
        risk_level = _quality_risk_level(quality.status)
        rationale = []

        if preprocessing:
            rationale.append("Applied {0} before OCR".format(", ".join(preprocessing)))
        else:
            rationale.append("Original image quality is acceptable for OCR")

        if quality.issues:
            rationale.append("Flagged quality risks: {0}".format(", ".join(quality.issues)))

        if self.review_policy == "review_all":
            rationale.append("User-selected policy requires review for every page")
        elif human_review_required:
            rationale.append("Quality gate requires human review before trusting final text")
        else:
            rationale.append("Semi-autonomous conversion can proceed with review option")

        return AgentDecision(
            autonomy_level="Semi-autonomous",
            preprocessing_applied=preprocessing,
            diagram_strategy="Detect visual regions, crop them, and place them near the matching text flow.",
            human_review_required=human_review_required,
            rationale=rationale,
            risk_level=risk_level,
            review_policy=self.review_policy,
            safety_controls=list(SAFETY_CONTROLS),
            external_tools=list(EXTERNAL_TOOLS),
        )

    def act(self, image_path: str, page_number: int = 1) -> PageConversionResult:
        trace = []
        warnings = []
        quality = self.observe(image_path)
        processed_path, preprocessing = prepare_image_for_ocr(image_path, quality, self.work_dir)
        decision = self.decide(quality, preprocessing)

        trace.append(
            "Observe: {0} is {1}x{2}px with brightness {3}, contrast {4}, sharpness {5}.".format(
                quality.filename,
                quality.width,
                quality.height,
                quality.brightness,
                quality.contrast,
                quality.blur_score,
            )
        )
        trace.append("Interpret: quality status is {0}.".format(quality.status))
        trace.append(
            "Decide: preprocessing={0}; review_required={1}; risk={2}.".format(
                ", ".join(preprocessing) or "none",
                "yes" if decision.human_review_required else "no",
                decision.risk_level,
            )
        )

        try:
            with open(processed_path, "rb") as image_file:
                extracted_text = extract_text_from_image(image_file.read(), os.path.basename(processed_path))
            trace.append("Act: OCR tool completed.")
        except Exception as error:
            extracted_text = "[OCR failed. Review the source image manually.]"
            message = "OCR tool failed: {0}".format(error)
            warnings.append(message)
            trace.append("Act: OCR tool failed and page was escalated for review.")
            _append_failure(decision, "OCR failed; human review is required before using this page.", "High")

        diagram_data = []
        try:
            diagram_output = extract_diagrams(processed_path)
            bounds = parse_diagram_bounds(diagram_output)
            trace.append("Act: diagram detector returned {0} candidate region(s).".format(len(bounds)))
            if bounds:
                prefix = "page_{0}_{1}".format(page_number, uuid.uuid4().hex[:8])
                diagram_paths = crop_diagrams(processed_path, bounds, self.diagram_dir, prefix)
                if len(diagram_paths) != len(bounds):
                    warnings.append("Some detected diagram regions could not be cropped.")
                    _append_failure(decision, "At least one diagram crop failed; check diagrams manually.", "Medium")

                for path, bound in zip(diagram_paths, bounds):
                    diagram_data.append(
                        {
                            "path": path,
                            "x_min": bound["x_min"],
                            "y_min": bound["y_min"],
                            "x_max": bound["x_max"],
                            "y_max": bound["y_max"],
                        }
                    )
        except Exception as error:
            message = "Diagram extraction failed: {0}".format(error)
            warnings.append(message)
            trace.append("Act: diagram extraction failed and page was escalated for manual diagram review.")
            _append_failure(decision, "Diagram extraction failed; verify visual content manually.", "Medium")

        trace.append(
            "Learn: page outcome captured with {0} diagram(s), {1} warning(s), and {2} review flag.".format(
                len(diagram_data),
                len(warnings),
                "a" if decision.human_review_required else "no",
            )
        )

        return PageConversionResult(
            source_path=image_path,
            processed_path=processed_path,
            extracted_text=extracted_text,
            diagram_data=diagram_data,
            quality=quality,
            decision=decision,
            trace=trace,
            warnings=warnings,
            status="completed_with_warnings" if warnings else "completed",
        )

    def learn(self, results: List[PageConversionResult]) -> Dict[str, object]:
        common_quality_issues = sorted(set(issue for item in results for issue in item.quality.issues))
        run_risk_level = _highest_risk(*[item.decision.risk_level for item in results])
        pages_requiring_review = [item.quality.filename for item in results if item.decision.human_review_required]

        return {
            "run_id": uuid.uuid4().hex[:12],
            "generated_at": datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ"),
            "pages_processed": len(results),
            "pages_requiring_review": len(pages_requiring_review),
            "review_pages": pages_requiring_review,
            "diagrams_detected": sum(len(item.diagram_data) for item in results),
            "tool_warnings": sum(len(item.warnings) for item in results),
            "common_quality_issues": common_quality_issues,
            "run_risk_level": run_risk_level,
            "review_policy": self.review_policy,
            "agent_type": "Goal-based semi-autonomous agent",
            "workflow": "Observe -> Interpret -> Decide -> Act -> Learn -> Human review",
            "memory_policy": "Short-term run memory only; no permanent user-note memory by default.",
            "human_checkpoint": "User reviews quality flags, OCR text, diagrams, and DOCX before final use.",
            "external_tools_used": list(EXTERNAL_TOOLS),
            "agentic_requirements": list(AGENTIC_REQUIREMENTS),
            "safety_controls": list(SAFETY_CONTROLS),
            "risk_register": list(RISK_REGISTER),
        }

    def run(self, image_paths: List[str], output_path: str) -> Tuple[str, str, List[PageConversionResult], Dict[str, object]]:
        results = []
        for page_number, image_path in enumerate(image_paths, start=1):
            results.append(self.act(image_path, page_number=page_number))

        memory = self.learn(results)
        docx_path = create_agentic_docx(results, output_path, memory)
        full_text = "\n\n".join(result.extracted_text for result in results)
        return full_text, docx_path, results, memory


def create_formatted_docx_inline(
    text: str, output_path: str, diagram_data: Optional[List[Dict[str, object]]] = None
) -> str:
    placeholder_quality = ImageQualityReport(
        filename="single_page",
        width=0,
        height=0,
        brightness=0,
        contrast=0,
        blur_score=0,
        status="Not assessed",
    )
    placeholder_decision = AgentDecision(
        autonomy_level="Semi-autonomous",
        preprocessing_applied=[],
        diagram_strategy="Inline placement",
        human_review_required=True,
        rationale=["Created from supplied OCR text without image-quality observation."],
    )
    result = PageConversionResult(
        source_path="",
        processed_path="",
        extracted_text=text,
        diagram_data=diagram_data or [],
        quality=placeholder_quality,
        decision=placeholder_decision,
    )
    return create_agentic_docx([result], output_path)


def process_image_to_docx_inline(
    image_path: str,
    output_path: str,
    review_policy: str = "quality_gate",
) -> Tuple[str, str]:
    agent = DocumentConversionAgent(review_policy=review_policy)
    extracted_text, docx_path, _, _ = agent.run([image_path], output_path)
    return extracted_text, docx_path


def process_images_to_docx(
    image_paths: List[str],
    output_path: str,
    review_policy: str = "quality_gate",
):
    agent = DocumentConversionAgent(review_policy=review_policy)
    return agent.run(image_paths, output_path)


def build_trace_payload(results: List[PageConversionResult], memory: Dict[str, object]) -> Dict[str, object]:
    return {
        "run_memory": memory,
        "pages": [
            {
                "source": os.path.basename(result.source_path),
                "processed": os.path.basename(result.processed_path),
                "status": result.status,
                "quality": asdict(result.quality),
                "decision": asdict(result.decision),
                "diagrams": [
                    {
                        "file": os.path.basename(str(diagram.get("path", ""))),
                        "x_min": diagram.get("x_min"),
                        "y_min": diagram.get("y_min"),
                        "x_max": diagram.get("x_max"),
                        "y_max": diagram.get("y_max"),
                    }
                    for diagram in result.diagram_data
                ],
                "warnings": list(result.warnings),
                "trace": list(result.trace),
            }
            for result in results
        ],
    }
