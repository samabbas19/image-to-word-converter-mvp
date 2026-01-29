from groq import Groq
import base64
import mimetypes
import os
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import cv2
from diagrams import extract_diagrams, parse_diagram_bounds, crop_diagrams

load_dotenv()
groq_key = os.getenv("GROK_API_KEY")
client = Groq(api_key=groq_key)

PROMPT = (
    "You are an OCR and transcription engine.\n\n"
    "Extract and reproduce ALL visible text and symbols from the provided image exactly as they appear.\n\n"
    "STRICT RULES:\n"
    "- Do NOT explain, summarize, reason, or interpret anything.\n"
    "- Do NOT add headings, comments, or descriptions.\n"
    "- Do NOT infer missing or unclear text.\n"
    "- Do NOT correct spelling, grammar, or formatting.\n"
    "- Do NOT reorganize or reformat content.\n"
    "- Do NOT describe diagrams — reproduce their text labels and structure using ASCII where needed.\n"
    "- Preserve original line breaks, spacing, indentation, arrows, boxes, bullet points, equations, and symbols.\n"
    "- If text is unclear or illegible, write: [illegible]\n"
    "- If a diagram contains text, include that text in its relative position.\n\n"
    "OUTPUT FORMAT:\n"
    "- Output ONLY the extracted content.\n"
    "- Plain text only.\n\n"
    "Any reasoning, explanation, or commentary is a violation of the task."
)

def image_bytes_to_data_url(image_bytes, filename="image.png"):
    """Convert image bytes to data URL for API"""
    mime_type, _ = mimetypes.guess_type(filename)
    if mime_type is None:
        mime_type = "image/png"
    
    encoded = base64.b64encode(image_bytes).decode("utf-8")
    return f"data:{mime_type};base64,{encoded}"

def extract_text_from_image(image_bytes, filename):
    """Extract text from image using Groq API"""
    image_url = image_bytes_to_data_url(image_bytes, filename)
    
    completion = client.chat.completions.create(
        model="meta-llama/llama-4-maverick-17b-128e-instruct",
        messages=[
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": PROMPT},
                    {"type": "image_url", "image_url": {"url": image_url}},
                ],
            }
        ],
        temperature=0,
        max_completion_tokens=2048,
    )
    
    return completion.choices[0].message.content

def detect_formatting(text):
    """
    Detect basic formatting patterns in extracted text
    Returns structured data with formatting hints
    """
    lines = text.split('\n')
    formatted_content = []
    
    for line in lines:
        line_data = {
            'text': line,
            'is_heading': False,
            'is_bullet': False,
            'is_numbered': False,
            'indent_level': 0,
            'alignment': 'left'
        }
        
        # Detect headings (ALL CAPS or lines with # prefix)
        if line.strip().isupper() and len(line.strip()) > 0 and len(line.strip()) < 100:
            line_data['is_heading'] = True
        
        # Detect bullet points
        if re.match(r'^\s*[•\-\*]\s+', line):
            line_data['is_bullet'] = True
            line_data['indent_level'] = len(line) - len(line.lstrip())
        
        # Detect numbered lists
        if re.match(r'^\s*\d+[\.\)]\s+', line):
            line_data['is_numbered'] = True
            line_data['indent_level'] = len(line) - len(line.lstrip())
        
        # Detect indentation
        if not line_data['is_bullet'] and not line_data['is_numbered']:
            line_data['indent_level'] = len(line) - len(line.lstrip())
        
        formatted_content.append(line_data)
    
    return formatted_content

def create_formatted_docx(text, output_path, diagram_paths=None):
    """
    Create a formatted DOCX file from extracted text and embed diagrams
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Detect formatting
    formatted_content = detect_formatting(text)
    
    for line_data in formatted_content:
        text_content = line_data['text'].strip()
        
        if not text_content:  # Empty line
            doc.add_paragraph()
            continue
        
        # Add paragraph
        p = doc.add_paragraph()
        
        # Apply heading style
        if line_data['is_heading']:
            p.style = 'Heading 1'
            run = p.add_run(text_content)
            run.bold = True
        
        # Apply bullet point
        elif line_data['is_bullet']:
            p.style = 'List Bullet'
            # Remove bullet character from text
            clean_text = re.sub(r'^[•\-\*]\s+', '', text_content)
            p.add_run(clean_text)
        
        # Apply numbered list
        elif line_data['is_numbered']:
            p.style = 'List Number'
            # Remove number from text
            clean_text = re.sub(r'^\d+[\.\\)]\s+', '', text_content)
            p.add_run(clean_text)
        
        # Regular paragraph
        else:
            p.add_run(text_content)
        
        # Apply indentation
        if line_data['indent_level'] > 0:
            p.paragraph_format.left_indent = Inches(line_data['indent_level'] * 0.05)
    
    # Add diagrams at the end with labels
    if diagram_paths and len(diagram_paths) > 0:
        doc.add_page_break()
        doc.add_heading('Extracted Diagrams', level=1)
        
        for idx, diagram_path in enumerate(diagram_paths, start=1):
            if os.path.exists(diagram_path):
                # Add diagram label
                p = doc.add_paragraph()
                run = p.add_run(f'Diagram {idx}:')
                run.bold = True
                
                # Add the image
                try:
                    doc.add_picture(diagram_path, width=Inches(5.5))
                except Exception as e:
                    # If image can't be added, note it
                    doc.add_paragraph(f'[Error loading diagram: {str(e)}]')
                
                # Add spacing
                doc.add_paragraph()
    
    # Save document
    doc.save(output_path)
    return output_path

def process_image_to_docx(image_path, output_path):
    """
    Main function to process image and create DOCX with diagrams
    """
    # Read image
    with open(image_path, 'rb') as f:
        image_bytes = f.read()
    
    # Extract text
    extracted_text = extract_text_from_image(image_bytes, os.path.basename(image_path))
    
    # Extract and crop diagrams
    diagram_paths = []
    try:
        print(f"\n=== Starting diagram extraction for {image_path} ===")
        
        # Get diagram bounds from LLM
        diagram_result = extract_diagrams(image_path)
        print(f"LLM Diagram Result:\n{diagram_result}\n")
        
        # Parse bounds
        bounds = parse_diagram_bounds(diagram_result)
        print(f"Parsed bounds: {bounds}")
        
        if bounds and len(bounds) > 0:
            print(f"Found {len(bounds)} diagram(s), cropping...")
            # Crop diagrams and save them
            diagram_paths = crop_diagrams(image_path, bounds)
            print(f"Successfully extracted {len(diagram_paths)} diagrams:")
            for path in diagram_paths:
                print(f"  - {path} (exists: {os.path.exists(path)})")
        else:
            print("No diagrams detected in the image")
            
    except Exception as e:
        print(f"ERROR: Could not extract diagrams: {str(e)}")
        import traceback
        traceback.print_exc()
        # Continue without diagrams if extraction fails
    
    # Create DOCX with text and diagrams
    print(f"\nCreating DOCX with {len(diagram_paths)} diagram(s)...")
    docx_path = create_formatted_docx(extracted_text, output_path, diagram_paths)
    print(f"DOCX created at: {docx_path}\n")
    
    return extracted_text, docx_path

